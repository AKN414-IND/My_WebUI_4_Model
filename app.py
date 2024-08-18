import streamlit as st
from PyPDF2 import PdfReader
from io import BytesIO
from PIL import Image
import docx
from langchain.chains import RetrievalQA
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.manager import CallbackManager
from langchain.llms import Ollama
from langchain.embeddings.ollama import OllamaEmbeddings
from langchain.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
import os

class Document:
    def __init__(self, text, metadata=None):
        self.page_content = text
        self.metadata = metadata if metadata else {}

LOCAL_MODEL = "llama3.1"
EMBEDDING = "nomic-embed-text"
llm = Ollama(base_url="http://localhost:11434", model=LOCAL_MODEL, verbose=True,
             callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]))

st.set_page_config(page_title="Chat-Based Query Answering", layout="wide")
st.title("Chat-Based Query Answering System")

with st.sidebar:
    st.header("Upload a File")
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "png", "jpg", "jpeg", "docx", "txt"])

if uploaded_file is not None:
    all_text = ""

    if uploaded_file.type == "application/pdf":
        pdf_file = BytesIO(uploaded_file.getvalue())
        reader = PdfReader(pdf_file)
        for page in range(len(reader.pages)):
            page_text = reader.pages[page].extract_text()
            if page_text:
                all_text += page_text + "\n"
    
    elif uploaded_file.type in ["image/png", "image/jpeg"]:
        image = Image.open(uploaded_file)
        st.image(image, caption="Uploaded Image", use_column_width=True)
        # OCR text extraction logic would be added here
    
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"

    elif uploaded_file.type == "text/plain":
        all_text = uploaded_file.read().decode("utf-8")
    
    with st.expander("Extracted Text", expanded=False):
        st.text_area("Text Content", all_text, height=200)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
    doc = Document(all_text)
    all_splits = text_splitter.split_documents([doc])

    persist_directory = 'data'
    if not os.path.exists(persist_directory):
        os.makedirs(persist_directory)

    vectorstore = Chroma.from_documents(
        documents=all_splits,
        embedding=OllamaEmbeddings(model=EMBEDDING),
        persist_directory=persist_directory
    )

    retriever = vectorstore.as_retriever()
    template = """You are a knowledgeable chatbot, here to help with questions of the user. Your tone should be professional and informative.

        Context: {context}
        History: {history}

        User: {question}
        Chatbot:
    """
    prompt = PromptTemplate(
        input_variables=["history", "context", "question"],
        template=template,
    )

    memory = ConversationBufferMemory(
        memory_key="history",
        return_messages=True,
        input_key="question"
    )

    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type='stuff',
        retriever=retriever,
        verbose=True,
        chain_type_kwargs={
            "verbose": True,
            "prompt": prompt,
            "memory": memory,
        }
    )

    st.write("---")
    st.subheader("Chat with the Document")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    with st.container():
        for i, msg in enumerate(st.session_state.messages):
            st.write(f"**{'User' if i % 2 == 0 else 'Chatbot'}:** {msg}")

    query = st.text_input("Enter your question:", value="", key="input_query", placeholder="Type your question here...")
    if st.button("Submit"):
        if query:
            st.session_state.messages.append(query)
            response = qa_chain.invoke({"query": query})
            st.session_state.messages.append(response['result'])
            st.experimental_rerun()


st.markdown(
    """
    <style>
    .stTextInput > div > div > input {
        font-size: 18px;
        padding: 10px;
        border-radius: 10px;
        border: 1px solid #ddd;
    }
    .stTextInput > div > div > input:focus {
        border: 1px solid #aaa;
    }
    .stContainer > div {
        margin-bottom: 20px;
        padding: 10px;
        border-radius: 10px;
        background-color: #f9f9f9;
    }
    .stContainer > div:nth-child(odd) {
        background-color: #e0e0e0;
    }
    .stTextArea {
        border-radius: 10px;
        border: 1px solid #ddd;
    }
    </style>
    """, unsafe_allow_html=True
)
